"""DOCX document processor using python-docx."""

from pathlib import Path
from typing import List
from uuid import uuid4

from src.models.document import Document, DocumentPage, FileType


class DocxProcessor:
    """Processor for DOCX documents."""

    def process(self, file_path: str) -> Document:
        """Process a DOCX file and extract text content.

        Extracts text from paragraphs, tables, headers, and footers.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Document model with extracted content.

        Raises:
            FileNotFoundError: If the file does not exist.
        """
        from docx import Document as DocxDocument

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = DocxDocument(file_path)
        sections = []

        # Extract header text
        for section in doc.sections:
            header = section.header
            if header:
                header_text = self._extract_paragraphs(header.paragraphs)
                if header_text:
                    sections.append(f"[Header]\n{header_text}")

        # Extract main body paragraphs
        body_text = self._extract_paragraphs(doc.paragraphs)
        if body_text:
            sections.append(body_text)

        # Extract tables
        table_text = self._extract_tables(doc.tables)
        if table_text:
            sections.append(f"[Tables]\n{table_text}")

        # Extract footer text
        for section in doc.sections:
            footer = section.footer
            if footer:
                footer_text = self._extract_paragraphs(footer.paragraphs)
                if footer_text:
                    sections.append(f"[Footer]\n{footer_text}")

        full_text = "\n\n".join(sections)

        page = DocumentPage(
            page_number=1,
            text=full_text,
            confidence=1.0,
        )

        return Document(
            id=uuid4(),
            filename=path.name,
            file_type=FileType.DOCX,
            content=full_text,
            raw_text=full_text,
            page_count=1,
            pages=[page],
            file_path=str(path.absolute()),
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )

    def _extract_paragraphs(self, paragraphs: List) -> str:
        """Extract text from a list of paragraph objects.

        Args:
            paragraphs: List of docx Paragraph objects.

        Returns:
            Combined paragraph text.
        """
        texts = []
        for para in paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)
        return "\n".join(texts)

    def _extract_tables(self, tables: List) -> str:
        """Extract text from table objects.

        Args:
            tables: List of docx Table objects.

        Returns:
            Formatted table text.
        """
        table_texts = []
        for table in tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_texts.append("\n".join(rows))
        return "\n\n".join(table_texts)
